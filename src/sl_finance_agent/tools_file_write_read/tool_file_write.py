"""LangChain tools for writing files to the local filesystem from within an agent.

Exposes two ``@tool``-decorated functions:

* :func:`tool_custom_file_write` -- writes plain-text content to a file using LangChain's
  built-in :class:`WriteFileTool`. The file path is validated to ensure its parent
  directory equals the project-wide file directory (``FILE_DIR``); any other location
  is rejected.
* :func:`tool_generate_word_doc` -- generates a Microsoft Word document (``.docx``) from
  a title and Markdown-formatted content using ``python-docx`` and saves it to a path
  validated by :class:`TargetPathValidation` (must start with ``FILE_DIR``). If the
  supplied ``store_path`` is not under ``FILE_DIR``, the file is saved to the default
  ``FILE_DIR`` root as a fallback.

Environment
-----------
- ``FILE_DIR`` (imported from :mod:`sl_finance_agent.common_utils`)
    The project-wide base directory that constrains where files may be written.
    For :func:`tool_custom_file_write` the parent directory of ``file_path`` must
    equal ``FILE_DIR``; for :func:`tool_generate_word_doc` the ``store_path`` must
    start with ``FILE_DIR`` (paths outside it fall back to ``FILE_DIR``).

Typical use
-----------
These tools are decorated with ``@tool`` (or ``@tool(args_schema=...)``) so they
can be registered directly with a LangChain agent::

    from tools_file_write_read.tool_file_write import (
        tool_custom_file_write,
        tool_generate_word_doc,
    )
"""

import os
from pathlib import Path
import re
import shutil
import time
from typing import Optional

from docx import Document
from langchain.tools import tool
from langchain_community.tools import WriteFileTool
from pydantic import BaseModel, Field, field_validator

from ..common_utils import get_logger, FILE_DIR
logger = get_logger(__name__)

@tool
def tool_custom_file_write(file_path:str, content: str) -> str:
    """Write plain-text content (e.g. ``.txt``, ``.json``, ``.md``, ``.csv``) to a file under ``FILE_DIR``.

    The ``file_path``'s parent directory must equal ``FILE_DIR``; otherwise the call is rejected.
    Empty ``content`` raises an error. The actual write is delegated to
    :class:`langchain_community.tools.WriteFileTool`.

    Args:
        file_path: Destination path (e.g. ``path/to/your/file.md``). Parent dir must be ``FILE_DIR``.
        content: The text content to write.

    Returns:
        The write result message on success, or an error message on failure.
    """
    # Initialize the tool
    # Ensure the root_dir is set to where you want the file to be created
    file_full_path = file_path
    try:
        dir_name = os.path.dirname(file_full_path)
        if dir_name == FILE_DIR:
            logger.info("Write file, file path is under current woring directory: %s", file_full_path)
        else:
            logger.error("⚠️ Write file: file path is not under the current file directory: %s", file_full_path)
            return f"""⚠️ Write file: file path is not under the current file directory: {file_full_path},
                                check the file path parameter"""
    except ValueError as e:
        logger.error("❌ Write file, Path is wrong: %s", file_full_path)
        return f"❌ Write file, {str(e)} Path is wrong: {file_full_path}, check the file path parameter"

    try:
        if content is None or len(content) == 0:
            logger.error("❌ Write file, content is empty")
            raise ValueError("❌ Write file, content is empty")

        write_tool = WriteFileTool(root_dir=file_full_path)

        # Write the file
        result = write_tool.run({
            "file_path": file_full_path,
            "text": content
        })
        
        logger.info("✅ Write file: %s, Result: %s", file_full_path, result)
        return str(result)

    except Exception as e:
        logger.error("❌ Write file error: %s", str(e))
        return f"Write file error: {str(e)}"

# --- MARKDOWN PARSER ---
def _add_markdown_content_to_docx(doc, content: str):
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip().replace("**", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip().replace("**", ""), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip().replace("**", ""), level=3)
        else:
            p = (
                doc.add_paragraph(style="List Bullet")
                if line.startswith(("- ", "* "))
                else doc.add_paragraph()
            )
            text = line[2:] if line.startswith(("- ", "* ")) else line
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

# Step 1: Define strict validation
class TargetPathValidation(BaseModel):
    """
    A strict validation for the file path of LLM call parameter
    """
    store_path: str = Field(description="The exact file path of LLM call")
    
    @field_validator('store_path')
    @classmethod
    def validate_file_path(cls, v: str) -> str:
        """
        Validate the file path of LLM call parameter before tool call
        """
        v = v.strip().lower()
        if not v.startswith(FILE_DIR):
            logger.error(" The file path is invalid: %s ... The correct should start with: %s", v, FILE_DIR)
            raise ValueError(
                f"Invalid file path '{v}'. The correct should start with '{FILE_DIR}'. Check the file path parameter and retry."
            )
        return v

# Step 2: Define tool with schema
@tool(args_schema=TargetPathValidation)
def tool_generate_word_doc(title: str, content: str, store_path:str) -> str:
    """
    Generate a Word document (.docx) containing the given title and content, and save to store_path.
    The content supports Markdown formatting (headings, bold text, bullet points).

    :param title: file title
    :param content: file content
    :paran store_path: file store path
    :return: file saved path
    """

    filename = f"{title.replace(' ', '_')}_{int(time.time())}.docx"
    logger.debug("Generating Word document: %s", filename)

    # 保存路径
    local_path_file = ""
    local_path = store_path
    if local_path.startswith(FILE_DIR):
        if os.path.isdir(local_path):
            local_path_file = os.path.join(local_path, filename)
        elif local_path.endswith(".docx"):
            local_path_file = local_path
        else:
            logger.error("❌ Invalid Word document generate PATH: %s, STOP and check the path parameter!", local_path)
            return f"Invalid document generate PATH: {local_path}, STOP and check the path parameter!"

        logger.info("📌 Specify Word document generate PATH: %s", local_path_file)
    else:
        os.makedirs(FILE_DIR, exist_ok=True)
        local_path_file = os.path.join(FILE_DIR, filename)
        logger.warning("❌ Invalid Word document generate PATH: %s, 📌 use the default file root path : %s ", store_path, local_path_file)

    try:
        doc = Document()
        doc.add_heading(title, 0)
        _add_markdown_content_to_docx(doc, content)
        doc.save(local_path_file)

        # return the file down load link
        logger.info("✅ Taget document generated successfully: %s", local_path_file)
        return "✅ Target document generated successfully: {local_path_file}"
    except Exception as e:
        logger.error("❌ Word generated Error: %s", e)
        return f"❌ Word generated Error: {str(e)}"

@tool
def copy_file_to_folder(
    source: str | Path,
    target_folder: str | Path,
    new_filename: Optional[str] = None,
    overwrite: bool = False,
) -> str:
    """Copy ``source`` into ``target_folder``.

    Args:
        source: Path to the file to copy.
        target_folder: Destination directory. Created (with parents) if missing.
        new_filename: Optional new name for the copied file. If ``None`` the
            original basename is kept.
        overwrite: If ``True``, an existing file with the same name will be
            replaced. Defaults to ``False`` to avoid accidental overwrites.

    Returns:
        The Name string of the newly created file.

    Raises:
        FileNotFoundError: If ``source`` does not exist or is not a file.
        NotADirectoryError: If ``target_folder`` exists but is not a directory.
        FileExistsError: If a file with the target name already exists and
            ``overwrite`` is ``False``.
    """
    src_path = Path(source)
    if not src_path.is_file():
        logger.warning(f"Source file not found: {src_path}")
        raise FileNotFoundError(f"Source file not found: {src_path}")

    dst_folder = Path(target_folder)
    dst_folder.mkdir(parents=True, exist_ok=True)
    if not dst_folder.is_dir():
        raise NotADirectoryError(f"Target path is not a directory: {dst_folder}")

    dst_path = dst_folder / (new_filename or src_path.name)
    if dst_path.exists() and not overwrite:
        raise FileExistsError(f"File already exists: {dst_path}")

    # ``shutil.copy2`` preserves metadata (mtime, atime) in addition to the
    # file contents, which is usually what you want for a simple copy.
    shutil.copy2(src_path, dst_path)
    return dst_path.name